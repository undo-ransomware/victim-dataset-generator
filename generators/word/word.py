from docx import Document
from docx.shared import Inches
import lorem
import random

def create_docx(num_of_files = 100):
	for n in range(int(num_of_files)):
		document = Document()
		document.add_heading(lorem.sentence(), 0)

		p = document.add_paragraph(lorem.sentence())
		p.add_run('bold').bold = True
		p.add_run(lorem.sentence())
		p.add_run('italic.').italic = True

		document.add_heading(lorem.sentence(), level=1)
		document.add_paragraph(lorem.sentence(), style='Intense Quote')

		document.add_paragraph(
			lorem.sentence(), style='List Bullet'
		)
		document.add_paragraph(
			lorem.sentence(), style='List Number'
		)

		for n in range(random.randint(1,101)):
			document.add_paragraph(
				lorem.sentence(), style='List Bullet'
			)

		records = (
			(lorem.sentence(), lorem.sentence(), lorem.sentence()),
			(lorem.sentence(), lorem.sentence(), lorem.sentence()),
			(lorem.sentence(), lorem.sentence(), lorem.sentence())
		)

		table = document.add_table(rows=1, cols=3)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = lorem.sentence()
		hdr_cells[1].text = lorem.sentence()
		hdr_cells[2].text = lorem.sentence()
		for qty, id, desc in records:
			row_cells = table.add_row().cells
			row_cells[0].text = str(qty)
			row_cells[1].text = id
			row_cells[2].text = desc

		document.add_page_break()

		filename = "{0}/{0}_{1:03d}.{2}".format("docx", n, 'docx')
		document.save(filename)
		
def main(args):
    create_docx(num_of_files = args[0])
    return 0
 
if __name__ == "__main__":
    import sys 
    status = main(sys.argv[1:])
    sys.exit(status)