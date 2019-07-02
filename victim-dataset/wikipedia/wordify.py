import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from summarizer import extract_keywords, summarize

class TextDocument:
	def __init__(self):
		self.document = Document()
	
	def init_full_text(self, text):
		self.keywords = extract_keywords(text)
	
	def make_title(self, title, subtitle, *subsubtitles):
		self.document.add_heading(title, 0)
		self.document.add_paragraph().add_run(subtitle).bold = True
		self.document.add_paragraph("\n".join (subsubtitles))
		# for spacing, according to classic bad formatting practice
		self.document.add_paragraph()

	def make_section(self, title, texts):
		summary = set(summarize(texts, self.keywords, 300))
		self.document.add_heading(title, 1)
		for bullet in summary:
			li = self.document.add_paragraph(style='List Bullet')
			li.add_run(bullet).italic = True
		for para in texts:
			if not para.isspace():
				self.document.add_paragraph(para)

	def make_image(self, filename, captions):
		with Image.open(filename) as img:
			width, height = img.size

		# usable width is 5.5" in the idiotic-inch based default template. not
		# making it any higher bcause then we get too many pages with just an
		# image.
		WIDTH, HEIGHT = 5.5, 5.5
		w = h = None
		if height / HEIGHT > width / WIDTH:
			h = Inches(HEIGHT)
		else:
			w = Inches(WIDTH)
		p = self.document.add_paragraph()
		p.add_run().add_picture(filename, width=w, height=h)
		p.add_run("\n")
		for cap in captions:
			p.add_run(cap).italic = True
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER

	def save(self, filename):
		self.document.save('media/' + filename + '.docx')
